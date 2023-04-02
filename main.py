from math import log
from statistics import variance, mean
from decimal import Decimal
import docx
import os
import datetime
import sys
import sqlite3
from PyQt6.QtSql import *
from datetime import datetime
from PyQt6 import uic, QtCore, QtGui, QtWidgets
from PyQt6.QtGui import QIcon
from MainForm import Ui_MainForm
from DeleteDialog import Ui_DeleteDialog
from EditDialog import Ui_EditDialog
from ErrorDialog import Ui_ErrorDialog
from AddDialog1 import Ui_AddDialog1
from AddDialog2 import Ui_AddDialog2
from Calendar import Ui_Calendar
from AnalysisDialog import Ui_AnalysisDialog

db_name = 'databases/db_tested (2).db'

def connect_db(db_name):
    '''Функция устанавливает соединение с базой данных'''
    db = QSqlDatabase.addDatabase('QSQLITE')
    db.setDatabaseName(db_name)
    if not db.open():
        print('Не удалось подключиться к базе')
        return False
    return db

def get_headers(tab_name):
    '''Функция возвращает заголовки по переданному названию таблицы'''
    con = sqlite3.connect(db_name)
    cur = con.cursor()
    sql = 'SELECT name FROM PRAGMA_TABLE_INFO("{}")'.format(tab_name)
    headers = tuple(header[0] for header in cur.execute(sql).fetchall())
    print('Заголовки:', cur.execute(sql).fetchall())
    cur.close()
    con.close()
    return headers

def get_all_act_isp():
    '''Функция возвращает кортежи: коды бумаги и даты погащения'''
    con = sqlite3.connect(db_name)
    cur = con.cursor()
    sql = 'SELECT * FROM {}'.format('Act_isp')
    regnum_data = tuple(line[0] for line in cur.execute(sql).fetchall())
    exec_date_data = tuple(line[1] for line in cur.execute(sql).fetchall())
    cur.close()
    con.close()
    return (regnum_data, exec_date_data)

def get_all_Activ():
    '''Функция возвращает кортежи: даты торгов, коды бумаги, цены, доходности'''
    con = sqlite3.connect(db_name)
    cur = con.cursor()
    sql = 'SELECT * FROM {}'.format('Activ')
    torg_date_data = tuple(line[0] for line in cur.execute(sql).fetchall())
    regnum_data = tuple(line[1] for line in cur.execute(sql).fetchall())
    price_data = tuple(line[2] for line in cur.execute(sql).fetchall())
    gain_end = tuple(line[3] for line in cur.execute(sql).fetchall())
    cur.close()
    con.close()
    return (torg_date_data, price_data, regnum_data, gain_end)

def get_rate_Activ():
    '''Функция возвращает столбец со ставкой'''
    con = sqlite3.connect(db_name)
    cur = con.cursor()
    sql = 'SELECT {} FROM {}'.format('ставка', 'Activ')
    rate = [line[0] for line in cur.execute(sql).fetchall()]
    cur.close()
    con.close()
    return rate

def get_count_rows(tab_name):
    '''Функция возвращает количество строк в переданной таблице'''
    con = sqlite3.connect(db_name)
    cur = con.cursor()
    sql = 'SELECT COUNT(*) FROM {}'.format(tab_name)
    count_rows = cur.execute(sql).fetchall()[0][0]
    cur.close()
    con.close()
    return count_rows

def create_filter_doc():
    '''Функция формирует документ с таблицей по текущему фильтру'''

    headers = [i for i in get_headers('Activ')] # получение списка заголовков
    # считывание введённых фильтров
    date_start = Ui_Main.date_start.text()
    date_end = Ui_Main.date_end.text()
    regnum = Ui_Main.regnum.text()
    price_start = Ui_Main.price_start.text()
    price_end = Ui_Main.price_end.text()
    gain_start = Ui_Main.gain_start.text()
    gain_end = Ui_Main.gain_end.text()

    con = sqlite3.connect(db_name)
    cur = con.cursor()
    condition = Activ.filter() # условие для фильтрации
    if condition != '':
        sql = 'SELECT * FROM {} WHERE {}'.format('Activ', condition) # таблица с фильтром
    else:
        sql = 'SELECT * FROM {}'.format('Activ') # таблица без фильтра
    filtered_table = cur.execute(sql).fetchall()

    if len(filtered_table) == 0:
        Ui_Main.textBrowser.setText('В таблице отсутствуют записи')
    else:
        # получение минимального и максимального значений доходности из отфильтрованной таблицы
        min_price = float(filtered_table[0][2].replace(',', '.'))
        max_price = min_price
        min_gain = float(filtered_table[0][3].replace(',', '.'))
        max_gain = min_gain
        for i in range(1, len(filtered_table)):

            price_i = float(filtered_table[i][2].replace(',', '.'))
            gain_i = float(filtered_table[i][3].replace(',', '.'))

            if price_i < min_price:
                min_price = price_i
            if price_i > max_price:
                max_price = price_i

            if gain_i < min_gain:
                min_gain = gain_i
            if gain_i > max_gain:
                max_gain = gain_i

        doc = docx.Document()
        paragraph = doc.add_paragraph()
        paragraph.alignment = 1
        paragraph.add_run('Данные торгов ценными бумагами').bold = True
        paragraph = doc.add_paragraph()
        paragraph.add_run('Фильтрация: ').bold = True
        applied_filters = []

        #внесение существующих фильтров в докукумент
        if date_start != '' and date_end != '':
            paragraph = doc.add_paragraph()
            paragraph.add_run('Дата торгов: ').italic = True
            paragraph.add_run(f'От {date_start} до {date_end}')
            applied_filters.append(headers[0].replace('_',' '))
        else:
            paragraph = doc.add_paragraph()
            paragraph.add_run('Дата торгов: ').italic = True
            paragraph.add_run(f'От {filtered_table[0][0]} до {filtered_table[-1][0]}')

        if regnum != '':
            paragraph = doc.add_paragraph()
            paragraph.add_run('Код бумаги: ').italic = True
            paragraph.add_run(regnum)
            applied_filters.append(headers[1].replace('_', ' '))
        else:
            doc.add_paragraph('Код бумаги: все')

        def two_demical_places(val):
            val = "{:0.2f}".format(float(val.replace(",", ".")))
            return val.replace('.',',')



        if price_start != '' and price_end != '':
            paragraph = doc.add_paragraph()
            paragraph.add_run('Цена: ').italic = True

            paragraph.add_run(f'От {two_demical_places(price_start)} до {two_demical_places(price_end)}')
            applied_filters.append(headers[2].replace('_', ' '))
        else:
            paragraph = doc.add_paragraph()
            paragraph.add_run('Цена: ').italic = True
            paragraph.add_run(f'От {two_demical_places(str(min_price))} до {two_demical_places(str(max_price))}')

        if gain_start != '' and gain_end != '':
            paragraph = doc.add_paragraph()
            paragraph.add_run('Доходность: ').italic = True
            paragraph.add_run(f'От {two_demical_places(gain_start)} до {two_demical_places(gain_end)}')
            applied_filters.append(headers[3].replace('_', ' '))
        else:
            paragraph = doc.add_paragraph()
            paragraph.add_run('Доходность: ').italic = True
            paragraph.add_run(f'От {two_demical_places(str(min_gain))} до {two_demical_places(str(max_gain))}')

        paragraph = doc.add_paragraph()
        paragraph.add_run('Примененные фильтры: ').italic = True
        paragraph.add_run(', '.join(map(str, applied_filters)))

        table = doc.add_table(rows=len(filtered_table)+1, cols=len(headers)) # добавление таблицы
        table.style = 'Table Grid' # применение стиля для таблицы

        # заполнение таблицы заголовками
        for i in range(len(headers)):
            table.rows[0].cells[i].text = headers[i].replace('_', ' ')

        # заполнение таблицы данными
        for i in range(1, len(filtered_table)+1):
            for j in range(len(headers)):
                table.rows[i].cells[j].text = filtered_table[i-1][j]

        doc.save('doc1.docx')


def Analysis_Dialog():
    '''Функция позволяет взаимодействовать с окном анализа'''

    global AnalysisDialog
    AnalysisDialog = QtWidgets.QDialog()
    ui_analysis = Ui_AnalysisDialog()
    ui_analysis.setupUi(AnalysisDialog)
    AnalysisDialog.show()

    def select_date_start():
        '''Функция позволяет выбрать дату начала интервала'''

        global Calendar
        Calendar = QtWidgets.QWidget()
        ui_cal = Ui_Calendar()
        ui_cal.setupUi(Calendar)
        Calendar.show()

        def add_and_close():
            global sel_date_start
            sel_date_start = ui_cal.calendarWidget.selectedDate()
            sel_date_start = sel_date_start.toString("dd.MM.yyyy")
            ui_analysis.line_date_start.setText(sel_date_start)
            Calendar.close()

        ui_cal.CalendarCheck.clicked.connect(add_and_close)

    def select_date_end():
        '''Функция позволяет выбрать дату конца интервала'''

        global Calendar
        Calendar = QtWidgets.QWidget()
        ui_cal = Ui_Calendar()
        ui_cal.setupUi(Calendar)
        Calendar.show()

        def add_and_close():
                global sel_date_end
                sel_date_end = ui_cal.calendarWidget.selectedDate()
                sel_date_end = sel_date_end.toString("dd.MM.yyyy")
                ui_analysis.line_date_end.setText(sel_date_end)
                Calendar.close()
        ui_cal.CalendarCheck.clicked.connect(add_and_close)

    def create_analysis_doc():
        '''Функция формирует документ с таблицей анализа'''
        date_start = ui_analysis.line_date_start.text() # дата начала интервала
        date_t = ui_analysis.line_date_end.text() # дата конца интервала
        if date_start != '' and date_t != '':
            doc = docx.Document()
            paragraph = doc.add_paragraph()
            paragraph.alignment = 1
            paragraph.add_run('Анализ статистических характеристик').bold = True

            paragraph = doc.add_paragraph()
            paragraph.add_run('Выбранный день торгов: ').italic = True
            paragraph.add_run(date_t)

            paragraph = doc.add_paragraph()
            paragraph.add_run('Календарная предыстория: ').italic = True
            paragraph.add_run(f'От {date_start} до {date_t}')

            table = doc.add_table(rows=len_regnum_day+1, cols=6)  # добавление таблицы
            table.style = 'Table Grid'  # применение стиля для таблицы

            # заполнение таблицы заголовками
            for i in range(6):
                table.rows[0].cells[i].text = ui_analysis.tableWidget.horizontalHeaderItem(i).text()

            # заполнение таблицы данными
            for i in range(1, len_regnum_day+1):
                for j in range(6):
                    table.rows[i].cells[j].text = ui_analysis.tableWidget.item(i-1, j).text()

            doc.save('doc2.docx')
        else:
            ui_analysis.textBrowser.setText('Расчет не произведен')


    def clear_table():
        '''Функция очищает таблицу и поля от текущих значений'''

        # очищаем каждую ячейку отдельно, так как функция clear стирает заголовки
        for i in range(58):
            ui_analysis.tableWidget.item(i, 0).setText('')
            ui_analysis.tableWidget.item(i, 1).setText('')
            ui_analysis.tableWidget.item(i, 2).setText('')
            ui_analysis.tableWidget.item(i, 3).setText('')
            ui_analysis.tableWidget.item(i, 4).setText('')
            ui_analysis.tableWidget.item(i, 5).setText('')
        ui_analysis.line_date_start.clear()
        ui_analysis.line_date_end.clear()
        ui_analysis.textBrowser.setText('')

    def calculates_characteristics():
        '''Функция рассчитывает среднее и дисперсию процентных ставок для каждого кода бумаги выбранного интревала'''

        # проверка даты
        if 'sel_date_start' not in globals() or ui_analysis.line_date_start.text() == '':
            ui_analysis.textBrowser.setText('Не выбрана дата начала предыстории')
            return

        if 'sel_date_end' not in globals() or ui_analysis.line_date_end.text() == '':
            ui_analysis.textBrowser.setText('Не выбрана дата конца предыстории')
            return

        date_start = datetime.strptime(sel_date_start, '%d.%m.%Y') # дата начала интервала
        date_end = datetime.strptime(sel_date_end, '%d.%m.%Y') # дата конца интервала

        if date_end < date_start:
            ui_analysis.textBrowser.setText('Начальная дата больше конечной')
            return
        ui_analysis.textBrowser.clear()

        # получение необходимых данных
        count_rows = get_count_rows('Activ')  # количество строк в таблице Activ
        torg_date_Activ = get_all_Activ()[0] # все даты торгов из таблицы Activ
        regnum_Activ = get_all_Activ()[2] # все коды бумаги из таблицы Activ
        rate = get_rate_Activ() # все значения ставки

        # получение уникальных дат торгов из всей таблицы
        set_torg_date = []
        for i in range(count_rows):
            if torg_date_Activ[i] not in set_torg_date:
                set_torg_date.append(torg_date_Activ[i])

        max_torg_date = max([datetime.strptime(i, '%d.%m.%Y') for i in set_torg_date])
        min_torg_date = min([datetime.strptime(i, '%d.%m.%Y') for i in set_torg_date])

        if date_end > max_torg_date:
            ui_analysis.textBrowser.setText('Выбранная дата превышает максимально допустимую для выбора дату')
            return

        if date_end < min_torg_date:
            ui_analysis.textBrowser.setText('Минимально допустимая для выбора дата превышает выбранную дату')
            return

        if sel_date_end not in set_torg_date:
            ui_analysis.textBrowser.setText('Выбранная дата не является датой торгов')
            return

        if abs(set_torg_date.index(sel_date_start) - set_torg_date.index(sel_date_end)) in [0,1]:
            ui_analysis.textBrowser.setText('Выбрана слишком малая предыстория')
            return


        previous_date_end = datetime.strptime(set_torg_date[set_torg_date.index(sel_date_end) - 1], '%d.%m.%Y') # t-1 день
        print(set_torg_date.index(sel_date_end))
        print(set_torg_date)
        print('t-1 день:', previous_date_end)

        def calc_char_diff_days(date_start, date_end, regnum_day):
            '''Функция рассчитывает среднее, дисперсию и размах процентных ставок для каждого кода бумаги выбранного интревала
            Параметры: дата начала интервала, дата конца интервала, уникальные коды даты конца интервала
            Возвращаемые значения: дисперсия, среднее, размах'''

            # инициализация дополнительных переменных
            regnum_interval = [] # коды бумаги из интервала
            rate_interval = [] # ставки из интервала
            one_regnum_rate = []  # ставки из интервала для одного кода бумаги
            mean_rate = [] # среднее ставки для каждого кода бумаги из интервала
            variance_rate = [] # дисперсия ставки для каждого кода из интервала
            rate_range = [] # размах ставки для каждого кода бумаги из интервала

            # заполнение списков значениями
            for i in range(count_rows):
                if date_start <= datetime.strptime(torg_date_Activ[i], '%d.%m.%Y') <= date_end:
                    rate_interval.append(float(rate[i].replace(',', '.')))
                    regnum_interval.append(regnum_Activ[i])
            print(regnum_interval,len(regnum_interval))
            for i in range(len(regnum_day)):
                for j in range(len(regnum_interval)):
                    if regnum_day[i] == regnum_interval[j]:
                        one_regnum_rate.append(rate_interval[j])

                # подсчет среднего и дисперсии (деление на ноль избегается)
                if len(one_regnum_rate) == 1:
                    variance_rate.append(0)
                else:
                    variance_rate.append('{:0.10f}'.format(variance(one_regnum_rate))) # 10 знаков после запятой у дисперсии

                mean_rate.append('{:0.6f}'.format(mean(one_regnum_rate))) # 6 знаков после запятой у среднего

                # подсчет размаха
                rate_range.append('{:0.5f}'.format(max(one_regnum_rate) - min(one_regnum_rate)))

                one_regnum_rate = []

            return (mean_rate, variance_rate, rate_range)

        regnum_day = [] # уникальные коды бумаги для дня t
        regnum_previous_day = [] # уникальные коды бумаги для дня t-1
        general_regnum_day = [] # уникальные коды бумаги общие для обоих дней

        for i in range(count_rows):

            if datetime.strptime(torg_date_Activ[i], '%d.%m.%Y') == date_end:
                if regnum_Activ[i] not in regnum_day:
                    regnum_day.append(regnum_Activ[i])

            if datetime.strptime(torg_date_Activ[i], '%d.%m.%Y') == previous_date_end:
                if regnum_Activ[i] not in regnum_previous_day:
                    regnum_previous_day.append(regnum_Activ[i])

        for i in range(len(regnum_day)):
            if regnum_day[i] in regnum_previous_day and regnum_day[i] not in general_regnum_day:
                general_regnum_day.append(regnum_day[i])

        # понадобится для формирования документов
        global len_regnum_day
        len_regnum_day = len(regnum_day)

        # получение кодов бумаги, средних, дисперсий и размахов для дня t
        mean_rate = calc_char_diff_days(date_start, date_end, general_regnum_day)[0]
        variance_rate = calc_char_diff_days(date_start, date_end, general_regnum_day)[1]
        rate_range = calc_char_diff_days(date_start, date_end, general_regnum_day)[2]

        # получение кодов бумаги, средних и дисперсий для дня t-1
        previous_mean_rate = calc_char_diff_days(date_start, previous_date_end, general_regnum_day)[0]
        previous_variance_rate= calc_char_diff_days(date_start, previous_date_end, general_regnum_day)[1]

        for i in range(len(general_regnum_day)):
            # печать среднего, дисперсии и размаха
            ui_analysis.tableWidget.item(i, 0).setText(general_regnum_day[i])
            ui_analysis.tableWidget.item(i, 1).setText(mean_rate[i])
            ui_analysis.tableWidget.item(i, 2).setText(variance_rate[i])
            ui_analysis.tableWidget.item(i, 5).setText(rate_range[i])

            # печать изменений среднего и дисперсии дня t по сравнению с t-1
            if mean_rate[i] > previous_mean_rate[i]:
                ui_analysis.tableWidget.item(i, 3).setText('растет')
            elif mean_rate[i] < previous_mean_rate[i]:
                ui_analysis.tableWidget.item(i, 3).setText('уменьшается')
            else:
                ui_analysis.tableWidget.item(i, 3).setText('не изменилось')

            if variance_rate[i] > previous_variance_rate[i]:
                ui_analysis.tableWidget.item(i, 4).setText('растет')
            elif variance_rate[i] < previous_variance_rate[i]:
                ui_analysis.tableWidget.item(i, 4).setText('уменьшается')
            else:
                ui_analysis.tableWidget.item(i, 4).setText('не изменилось')


    # кнопки взаимодействия с окном
    ui_analysis.date_button_from.clicked.connect(select_date_start)
    ui_analysis.date_button_to.clicked.connect(select_date_end)
    ui_analysis.calc_button.clicked.connect(calculates_characteristics)
    ui_analysis.clear_button.clicked.connect(clear_table)
    ui_analysis.save_button.clicked.connect(create_analysis_doc)
    ui_analysis.exit_button.clicked.connect(AnalysisDialog.close)


def get_rate():
    '''
    Функция расчитывает и выводит в консоль список состоящий из процентных ставкок для каждого торгового дня
    Формула для расчета: xk(i) = ln{ Fk(i) / 100)}/( Tk - ti ), где
    Fk - текущая цена ценной бумаги
    Tk - дата погашения ценной бумаги
    ti - порядковый номер торгового дня
    '''

    def sql_update_query():
        sql = 'UPDATE {} SET {} = "{}" WHERE {} = "{}" AND {} = "{}" AND {} = "{}" AND {} = "{}" '. \
            format('Activ', headers[4], str(xk[i]).replace('.', ','),
                   headers[0], torg_date_Activ[i],
                   headers[1], regnum_Activ[i],
                   headers[2], price[i],
                   headers[3], gain_end[i])

        query.prepare(sql)
        flag = query.exec()
        print(flag)

    xk, Tk, t, res = [], [], [], [] # переменные для расчета

    count_rows = get_count_rows('Activ') # количество строк в таблице Activ
    headers = get_headers('Activ') # заголовки таблицы Activ

    regnum_act_isp = get_all_act_isp()[0] #все коды из таблицы act_isp
    exec_date = get_all_act_isp()[1] #все даты погашения из таблицы act_isp

    torg_date_Activ = get_all_Activ()[0] #все даты торгов из таблицы Activ
    price = get_all_Activ()[1]  # кортеж из цен для каждого торгового дня
    regnum_Activ = get_all_Activ()[2] #все коды из таблицы Activ
    gain_end = get_all_Activ()[3]  # все доходности из таблицы Activ

    # создание списка из дат погашения по соответствию кодов двух таблиц
    for reg_Activ in regnum_Activ:
        for index, reg_act_isp in enumerate(regnum_act_isp):
            if reg_Activ == reg_act_isp:
                res.append(exec_date[index])

    t = torg_date_Activ
    Tk = res
    Fk = [float(i.replace(',', '.')) for i in price]  # приведение цены к численному типу

    print(Tk[:7], t[:7], Fk[:7])
    # расчет ставки
    for i in range(count_rows):
        if Tk[i] == t[i]:
            xk.append(0)
        else:
                delta = (datetime.strptime(Tk[i], "%d.%m.%Y") - datetime.strptime(t[i], "%d.%m.%Y")).days
                xk.append('{:0.4f}'.format(log(Fk[i] / 100) / delta))

    print(xk)
    query = QSqlQuery()

    # добавление нового столбца со значениями ставки
    if 'ставка' not in headers:
        sql = 'ALTER TABLE {} ADD "ставка" TEXT NOT NULL DEFAULT ""'.format('Activ')
        query.prepare(sql)
        flag = query.exec()
        print(flag)
        Activ.select()

        headers = get_headers('Activ')  # заголовки таблицы Activ, включая ставку
        print(headers)
        for i in range(count_rows):
            sql_update_query()
        Activ.select()

    irrelevant_rate = get_rate_Activ() #ставка, которая находится в данный момент в таблице (может быть неактуальной)
    price = get_all_Activ()[1]  # кортеж из цен для каждого торгового дня

    # обновление значений ставки для добавленных строк
    while '' in irrelevant_rate:
        i = irrelevant_rate.index('')
        irrelevant_rate[i] = xk[i]
        sql_update_query()
    Activ.select()

    # обновление значения ставки для отредактированных строк
    for i in range(count_rows):
            if irrelevant_rate[i] != str(xk[i]).replace('.',','):
                sql_update_query()
                irrelevant_rate[i] = str(xk[i]).replace('.',',')
    Activ.select()


def reset_filter():
    '''Функция сбрасывает текущий фильтр'''
    Activ.setFilter('')
    Ui_Main.date_start.setText('')
    Ui_Main.date_end.setText('')
    Ui_Main.regnum.setText('')
    Ui_Main.price_start.setText('')
    Ui_Main.price_end.setText('')
    Ui_Main.gain_start.setText('')
    Ui_Main.gain_end.setText('')


def set_filter():
    ''''Функция фильтрует по заданным параметрам'''

    # считывание введённых фильтров
    date_start = Ui_Main.date_start.text()
    date_end = Ui_Main.date_end.text()
    regnum = Ui_Main.regnum.text()
    price_start = Ui_Main.price_start.text()
    price_end = Ui_Main.price_end.text()
    gain_start = Ui_Main.gain_start.text()
    gain_end = Ui_Main.gain_end.text()

    # получение необходимых данных
    count_rows = get_count_rows('Activ') # количество строк в таблице Activ
    torg_date_Activ = get_all_Activ()[0] # все даты торгов из таблицы Activ
    price_Activ = get_all_Activ()[1] # все цены из таблицы Activ
    regnum_Activ = get_all_Activ()[2] # все коды бумаги из таблицы Activ
    gain_end_Activ = get_all_Activ()[3] # все доходности бумаги из таблицы Activ

    def handle_date(date_start, date_end):
        '''Функция проверяет правильность введенных дат'''
        try:
            date_start = datetime.strptime(date_start, '%d.%m.%Y')
            date_end = datetime.strptime(date_end, '%d.%m.%Y')
        except:
            Ui_Main.textBrowser.setText('Введена дата неверного формата')
            return False

        if date_start > date_end:
            Ui_Main.textBrowser.setText('Начальная дата больше конечной')
            return False
        if date_start > max([datetime.strptime(i, '%d.%m.%Y') for i in torg_date_Activ]):
            Ui_Main.textBrowser.setText('Начальная дата торгов больше максимально возможной')
            return False
        if date_start < min([datetime.strptime(i, '%d.%m.%Y') for i in torg_date_Activ]):
            Ui_Main.textBrowser.setText('Конечная дата торгов меньше минимально возможной')
            return False
        Ui_Main.textBrowser.clear()
        return True

    def handle_regnum():
        '''Функция проверяет правильность введенного кода'''
        if regnum[:2] != 'SU' or regnum[-4:] != 'RMFS' or regnum[2:7].isdigit() != True:
            Ui_Main.textBrowser.setText('Введено недопустимое значение кода')
            return False
        if regnum not in regnum_Activ:
            Ui_Main.textBrowser.setText('Введенный код не найден')
            return False
        Ui_Main.textBrowser.clear()
        return True

    def handle_price():
        '''Функция проверяет правильность введенных цен'''
        try:
            if float(price_start) < 0 or float(price_end) < 0:
                Ui_Main.textBrowser.setText('Введено отрицательное значение')
                return False
        except ValueError:
            Ui_Main.textBrowser.setText('Введено значение неверного типа')
            return False
        if float(price_start) > float(price_end):
            Ui_Main.textBrowser.setText('Начальное значение цены больше конечного')
            return False
        if float(price_start) > max([float(i.replace(',','.')) for i in price_Activ]):
            Ui_Main.textBrowser.setText('Начальное значение цены больше максимально возможного')
            return False
        if float(price_end) < min([float(i.replace(',','.')) for i in price_Activ]):
            Ui_Main.textBrowser.setText('Конечное значение цены меньше минимально возможного')
            return False
        Ui_Main.textBrowser.clear()
        return True

    def handle_gain():
        '''Функция проверяет правильность введенных доходностей'''
        try:
            if float(gain_start) < 0 or float(gain_end) < 0:
                Ui_Main.textBrowser.setText('Введено отрицательное значение')
                return False
        except ValueError:
            Ui_Main.textBrowser.setText('Введено значение неверного типа')
            return False
        if float(gain_start) > float(gain_end):
            Ui_Main.textBrowser.setText('Начальное значение доходности больше конечного')
            return False
        if float(gain_start) > max([float(i.replace(',','.')) for i in gain_end_Activ]):
            Ui_Main.textBrowser.setText('Начальное значение доходности больше максимально возможного')
            return False
        if float(gain_end) < min([float(i.replace(',','.')) for i in gain_end_Activ]):
            Ui_Main.textBrowser.setText('Конечное значение доходности меньше минимально возможного')
            return False
        Ui_Main.textBrowser.clear()
        return True

    def get_date_between():
        '''Функция возвращает кортеж из уникальных дат выбранного диапазона'''
        date_between = tuple()
        for i in range(count_rows):
            if datetime.strptime(date_start, '%d.%m.%Y') <= datetime.strptime(torg_date_Activ[i], '%d.%m.%Y') \
                    <= datetime.strptime(date_end, '%d.%m.%Y') and torg_date_Activ[i] not in date_between:
                date_between += (torg_date_Activ[i],)
        if len(date_between) == 1:
            date_between *= 2
        return date_between

    def get_price_between():
        '''Функция возвращает кортеж из цен выбранного диапазона'''
        price_between = tuple()
        for i in range(count_rows):
            if float(price_start.replace(',', '.')) <= float(price_Activ[i].replace(',', '.')) <= float(price_end.replace(',', '.')):
                price_between += (price_Activ[i],)
        if len(price_between) == 1:
            price_between *= 2
        return price_between

    def get_gain_between():
        '''Функция возвращает кортеж из доходностей выбранного диапазона'''
        gain_between = tuple()
        for i in range(count_rows):
            if float(gain_start.replace(',', '.')) <= float(gain_end_Activ[i].replace(',', '.')) <= float(gain_end.replace(',', '.')):
                gain_between += (gain_end_Activ[i],)
        if len(gain_between) == 1:
            gain_between *= 2
        return gain_between

    # фильтрация для всех возможных вариантов
    if  date_start == '' and date_end == '' and regnum == '' and price_start == '' and price_end == '' and gain_start != '' and gain_end != '':
        if handle_gain() == True:
            sql = 'текущая_доходность in {}'.format(get_gain_between())
            Activ.setFilter(sql)
            print(sql)

    if  date_start == '' and date_end == '' and regnum == '' and price_start != '' and price_end != '' and gain_start == '' and gain_end == '':
        if handle_price() == True:
            sql = 'цена in {}'.format(get_price_between())
            Activ.setFilter(sql)
            print(sql)

    if date_start == '' and date_end == '' and regnum == '' and price_start != '' and price_end != '' and gain_start != '' and gain_end != '':
        if handle_price() == True and handle_gain() == True:
            sql = 'цена in {} and текущая_доходность in {}'.format(get_price_between(), get_gain_between())
            Activ.setFilter(sql)
            print(sql)

    if date_start == '' and date_end == '' and regnum != '' and price_start == '' and price_end == '' and gain_start == '' and gain_end == '':
        if handle_regnum() == True:
            sql = 'код_бумаги = "{}"'.format(regnum)
            Activ.setFilter(sql)
            print(Activ.filter())
            print(sql)

    if date_start == '' and date_end == '' and regnum != '' and price_start == '' and price_end == '' and gain_start != '' and gain_end != '':
        if handle_regnum() == True and handle_gain() == True:
            sql = 'код_бумаги = "{}" and текущая_доходность in {}'.format(regnum, get_gain_between())
            Activ.setFilter(sql)
            print(sql)

    if date_start == '' and date_end == '' and regnum != '' and price_start != '' and price_end != '' and gain_start == '' and gain_end == '':
        if handle_regnum() == True and handle_price() == True:
            sql = 'код_бумаги = "{}" and цена in {}'.format(regnum, get_price_between())
            Activ.setFilter(sql)
            print(sql)

    if date_start == '' and date_end == '' and regnum != '' and price_start != '' and price_end != '' and gain_start != '' and gain_end != '':
        if handle_regnum() == True and handle_price() == True and handle_gain() == True:
            sql = 'код_бумаги = "{}" and цена in {} and текущая_доходность in {}'.format(regnum, get_price_between(), get_gain_between())
            Activ.setFilter(sql)
            print(sql)

    if date_start != '' and date_end != '' and regnum == '' and price_start == '' and price_end == '' and gain_start == '' and gain_end == '':
        if handle_date(date_start, date_end) == True:

            sql = 'дата_торгов in {}'.format(get_date_between())
            Activ.setFilter(sql)
            print(sql)

    if date_start != '' and date_end != '' and regnum == '' and price_start == '' and price_end == '' and gain_start != '' and gain_end != '':
        if handle_date(date_start, date_end) == True and handle_gain() == True:
            sql = 'дата_торгов in {} and текущая_доходность in {}'.format(get_date_between(), get_gain_between())
            Activ.setFilter(sql)
            print(sql)

    if date_start != '' and date_end != '' and regnum == '' and price_start != '' and price_end != '' and gain_start == '' and gain_end == '':
        if handle_date(date_start, date_end) == True and handle_price() == True:
            sql = 'дата_торгов in {} and цена in {}'.format(get_date_between(), get_price_between())
            Activ.setFilter(sql)
            print(sql)

    if date_start != '' and date_end != '' and regnum == '' and price_start != '' and price_end != '' and gain_start != '' and gain_end != '':
        if handle_date(date_start, date_end) == True and handle_regnum() == True and handle_price() == True:
            sql = 'дата_торгов in {} and цена in {} and текущая_доходность in {}'.format(get_date_between(), get_price_between(), get_gain_between())
            Activ.setFilter(sql)
            print(sql)

    if date_start != '' and date_end != '' and regnum != '' and price_start == '' and price_end == '' and gain_start == '' and gain_end == '':
        if handle_date(date_start, date_end) == True and handle_regnum() == True:
            sql = 'дата_торгов in {} and код_бумаги = "{}"'.format(get_date_between(), regnum)
            Activ.setFilter(sql)
            print(sql)

    if date_start != '' and date_end != '' and regnum != '' and price_start == '' and price_end == '' and gain_start != '' and gain_end != '':
        if handle_date(date_start, date_end) == True and handle_regnum() == True and handle_gain() == True:
            sql = 'дата_торгов in {} and код_бумаги = "{}" and текущая_доходность in {}'.format(get_date_between(), regnum, get_gain_between())
            Activ.setFilter(sql)
            print(sql)

    if date_start != '' and date_end != '' and regnum != '' and price_start != '' and price_end != '' and gain_start == '' and gain_end == '':
        if handle_date(date_start, date_end) == True and handle_regnum() == True and handle_price():
            sql = 'дата_торгов in {} and код_бумаги = "{}" and цена in {}'.format(get_date_between(), regnum, get_price_between())
            Activ.setFilter(sql)
            print(sql)

    if date_start != '' and date_end != '' and regnum != '' and price_start != '' and price_end != '' and gain_start != '' and gain_end != '':
        if handle_date(date_start, date_end) == True and handle_regnum() == True and handle_price() and handle_gain() == True:
            sql = 'дата_торгов in {} and код_бумаги = "{}" and цена in {} and текущая_доходность in {}'\
                .format(get_date_between(), regnum, get_price_between(), get_gain_between())
            Activ.setFilter(sql)
            print(sql)


def Add_Dialog():
    '''Функция открывает диалог добавления для выбранной таблицы'''
    tab_name = Ui_Main.tabWidget.tabText(Ui_Main.tabWidget.currentIndex())  # имя открытой в данный момент таблицы
    if tab_name == 'Activ':
        Add_Dialog_Activ()
    elif tab_name == 'Act_isp':
        Add_Dialog_act_isp()


def Add_Dialog_Activ():
    '''Функция позволяет взаимодействовать с окном добавления'''

    global AddDialog1
    AddDialog1 = QtWidgets.QWidget()
    ui_add_tab1 = Ui_AddDialog1()
    ui_add_tab1.setupUi(AddDialog1)
    AddDialog1.show()

    # обновляем выпадающий список кодами из act_isp
    regnum_fresh = sorted(list(get_all_act_isp()[0])) # свежие коды бумаги
    ui_add_tab1.comboBox_regnum.clear()
    ui_add_tab1.comboBox_regnum.addItems(regnum_fresh)

    def add_date():
        '''Функция позволяет выбрать дату из календаря'''

        global Calendar
        Calendar = QtWidgets.QWidget()
        ui_cal = Ui_Calendar()
        ui_cal.setupUi(Calendar)
        Calendar.show()

        def add_and_close():
            global added_date  # выбранная дата из календаря
            added_date = ui_cal.calendarWidget.selectedDate()
            added_date = added_date.toString("dd.MM.yyyy")
            ui_add_tab1.line_date.setText(added_date)
            Calendar.close()

        ui_cal.CalendarCheck.clicked.connect(add_and_close)

    def sql_insert_query(new_values, headers):
        '''Функция делает sql запрос на обновление таблицы
        Параметры: новая запись, заголовки'''
        query = QSqlQuery()
        sql = 'INSERT INTO {} {} VALUES {}'.format('Activ', headers, new_values)
        query.prepare(sql)
        flag = query.exec()
        print(flag)
        Activ.select()
        get_rate() # вызов функции для расчета значения ставки
        AddDialog1.close()

    def add_row_Activ():

        new_values = tuple() #кортеж для всех значений строки
        price = ui_add_tab1.line_price.text() # получение текущего значения цены
        gain_end = ui_add_tab1.line_gain_end.text() # получение текущего значения доходности

        # проверка на выбор даты
        if 'added_date' not in globals() or ui_add_tab1.line_date.text() == '':
            ui_add_tab1.textBrowser.setText('Не выбрана дата торгов')
            return

        new_values += (added_date,) # добавление к общему списку выбранной даты
        new_values += (ui_add_tab1.comboBox_regnum.currentText(),)  # добавление к общему списку выбранного кода

        # проверка на ввод текущей цены
        if price == '':
            ui_add_tab1.textBrowser.setText('Не указано значение текущей цены')
            return

        # проверка на правильность ввода текущей цены
        try:
            if float(price.replace(',','.')) < 0:
                ui_add_tab1.textBrowser.setText('Введено отрицательное значение')
                return
            price = '{:0.2f}'.format(float(price.replace(',', '.'))) # приводим к строчному формату и окргуляем до 2 знаков
            new_values += (price.replace('.',','),)
        except ValueError:
            ui_add_tab1.textBrowser.setText('Введено значение неверного типа')
            return

        # проверка на ввод текущей доходности
        if gain_end == '':
            ui_add_tab1.textBrowser.setText('Не указано значение текущей доходности')
            return

        # проверка на правильность ввода доходности
        try:
            if float(gain_end.replace(',','.')) <= 0:
                ui_add_tab1.textBrowser.setText('Введено отрицательное значение')
                return
            gain_end = '{:0.2f}'.format(float(gain_end.replace(',', '.')))  # приводим к строчному формату и окргуляем до 2 знаков
            new_values += (gain_end.replace('.',','),)
        except ValueError:
            ui_add_tab1.textBrowser.setText('Введено значение неверного типа')
            return

        headers = get_headers('Activ')
        if 'ставка' in headers:
            new_values += ('',)

        print('Новая запись:', new_values)
        ui_add_tab1.comboBox_regnum.addItem(new_values[0])
        sql_insert_query(new_values, headers)

    ui_add_tab1.date_button.clicked.connect(add_date)
    ui_add_tab1.button_ok.clicked.connect(add_row_Activ)
    ui_add_tab1.button_cancel.clicked.connect(AddDialog1.close)


def Add_Dialog_act_isp():
    '''Функция добавляет строку в таблицу Act_isp'''

    global AddDialog2
    AddDialog2 = QtWidgets.QWidget()
    ui_add_tab2 = Ui_AddDialog2()
    ui_add_tab2.setupUi(AddDialog2)
    AddDialog2.show()

    def add_date():
        '''Функция позволяет выбрать дату из календаря'''

        global Calendar
        Calendar = QtWidgets.QWidget()
        ui_cal = Ui_Calendar()
        ui_cal.setupUi(Calendar)
        Calendar.show()

        def add_and_close():
            global added_date  # выбранная дата из календаря
            added_date = ui_cal.calendarWidget.selectedDate()
            added_date = added_date.toString("dd.MM.yyyy")
            ui_add_tab2.line_date.setText(added_date)
            Calendar.close()

        ui_cal.CalendarCheck.clicked.connect(add_and_close)

    def sql_insert_query(new_values):
        '''Функция делает sql запрос на обновление таблицы
        Параметры: новая запись'''
        headers = get_headers('Act_isp')
        query = QSqlQuery()
        sql = 'INSERT INTO {} {} VALUES {}'.format('Act_isp', headers, new_values)
        query.prepare(sql)
        flag = query.exec()
        print(flag)
        act_isp.select()

        AddDialog2.close()

    def add_row_act_isp():

        new_values = tuple() #кортеж для всех значений строки
        regnum = ui_add_tab2.line_regnum.text() # введенный код бумаги

        # проверка кода бумаги
        if regnum == '':
            ui_add_tab2.textBrowser.setText('Не указан код бумаги')
            return
        if regnum.isdigit() != True:
            ui_add_tab2.textBrowser.setText('Введенный код неверного типа')
            return
        if len(regnum) != 5:
            ui_add_tab2.textBrowser.setText('Недостаточно сиволов кода')
            return
        if 'SU' + regnum + 'RMFS' in get_all_act_isp()[0]:
            ui_add_tab2.textBrowser.setText('Введенный код уже существует')
            return
        new_values += ('SU' + regnum + 'RMFS',)

        # проверка на выбор даты
        if 'added_date' not in globals() or ui_add_tab2.line_date.text() == '':
            ui_add_tab2.textBrowser.setText('Не выбрана дата торгов')
            return
        if added_date in get_all_act_isp()[1]:
            ui_add_tab2.textBrowser.setText('Введеная дата уже существует')
            return
        new_values += (added_date,)

        print('Новая запись:', new_values)
        sql_insert_query(new_values)

    ui_add_tab2.date_button.clicked.connect(add_date)
    ui_add_tab2.button_ok.clicked.connect(add_row_act_isp)
    ui_add_tab2.button_cancel.clicked.connect(AddDialog2.close)


def Edit_Dialog():
    '''Функция позволяет взаимодействовать с окном редактирования'''

    def ErrorDialog():
        '''Функция открывает окно с оповещением о том, что запись не выбрана'''
        global ErrorDialog
        ErrorDialog = QtWidgets.QDialog()
        ui_error = Ui_ErrorDialog()
        ui_error.setupUi(ErrorDialog)
        ErrorDialog.show()

        # закрытие диалогового окна при нажатии на кнопку
        ui_error.button_okk.clicked.connect(ErrorDialog.close)

    def sql_update_query(headers, price_cell, gain_end_cell):
        '''Функция делает sql запрос на обновление таблицы
        Параметры: заголовки, значение цены, значение доходности'''

        query = QSqlQuery()
        sql = 'UPDATE {} SET {} = "{}", {} = "{}" WHERE {} = "{}" AND {} = "{}" AND {} = "{}" AND {} = "{}" '. \
            format('Activ', headers[2], price_cell.replace('.', ','),
                   headers[3], gain_end_cell.replace('.', ','),
                   headers[0], list_values[0],
                   headers[1], list_values[1],
                   headers[2], list_values[2],
                   headers[3], list_values[3])

        query.prepare(sql)
        flag = query.exec()
        print(flag)
        Activ.select()
        get_rate() # вызов функции для обновления ставки
        EditDialog.close()

    def edit_row():

        # создание переменных с соответсвующими текущими значениями, записанных в QLineEdit
        price_cell = ui_edit.line_price.text()
        gain_end_cell = ui_edit.line_gain_end.text()
        # получение заголовков таблицы
        headers = get_headers('Activ')

        # проверка на правильность ввода текущей цены
        try:
            if float(price_cell.replace(',','.')) <= 0:
                ui_edit.textBrowser.setText('Введено неверное значение цены')
                return
        except ValueError:
            ui_edit.textBrowser.setText('Введено неверное значение цены')
            return

        # проверка на правильность ввода доходности
        try:
            if float(gain_end_cell.replace(',', '.')) <= 0:
                ui_edit.textBrowser.setText('Введено неверное значение доходности')
                return
        except ValueError:
            ui_edit.textBrowser.setText('Введено неверное значение доходности')
            return

        # проверка на наличие изменений
        if  [price_cell, gain_end_cell] == list_values[-2:]:
            ui_edit.textBrowser.setText('Изменения не были внесены')
            return

        # приводим к строчному формату и окргуляем до 2 знаков
        price_cell = '{:0.2f}'.format(float(price_cell.replace(',', '.')))
        gain_end_cell = '{:0.2f}'.format(float(gain_end_cell.replace(',', '.')))

        print('Измененные цена и доходность:', price_cell, gain_end_cell)

        sql_update_query(headers, price_cell, gain_end_cell)

    list_values = get_row() # значения выбранной записи из таблицы

    # если запись не выбрана, вызываем окно ошибки
    if list_values == -1:
        ErrorDialog()
    else:

        global EditDialog
        EditDialog = QtWidgets.QDialog()
        ui_edit = Ui_EditDialog()
        ui_edit.setupUi(EditDialog)
        EditDialog.show()

        # отображение значений из таблицы в QLineEdit
        ui_edit.line_torg_date.setText(list_values[0])
        ui_edit.line_regnum.setText(list_values[1])
        ui_edit.line_price.setText(list_values[2])
        ui_edit.line_gain_end.setText(list_values[3])

        ui_edit.button_ok.clicked.connect(edit_row)
        ui_edit.button_cancel.clicked.connect(EditDialog.close)


def Delete_Dialog():
    '''Функция позволяет взаимодейстовать с окном удаления'''

    def remove_row():
        '''Функция удаляет выбранные строки'''
        selected = Ui_Main.tableView_Activ.selectedIndexes()
        index_set = set(index.row() for index in selected)  # множество индексов выбранных строк
        for index in index_set:
            Activ.removeRow(index)
        Activ.select()

        DeleteDialog.close()

    global DeleteDialog
    DeleteDialog = QtWidgets.QDialog()
    ui_delete = Ui_DeleteDialog()
    ui_delete.setupUi(DeleteDialog)
    DeleteDialog.show()

    ui_delete.No_button.clicked.connect(DeleteDialog.close)
    ui_delete.Yes_button.clicked.connect(remove_row)


def get_row():
    '''Функция возвращает список со значениями текущей строки,
    если строка не выбрана, то возвращает -1'''
    selected = Ui_Main.tableView_Activ.currentIndex().row()
    if selected == -1:
        return selected
    else:
        list_values = [Ui_Main.tableView_Activ.model().index(selected, i).data() for i in range(4)]
        return list_values

def get_Activ():
    '''Функция создает обьъект таблицу Activ'''
    global Activ
    Activ = QSqlTableModel()  # создаем объект
    Activ.setTable('Activ')  # привязываем к объекту таблицу
    Activ.select()  # выбираем все строки из таблицы
    return Activ

def get_act_isp():
    '''Функция создает объект таблицу act_isp'''
    global act_isp
    act_isp = QSqlTableModel()
    act_isp.setTable('act_isp')
    act_isp.select()
    return act_isp

if __name__ == "__main__":

    #Проверка подключения к базе данных
    if not connect_db(db_name):
        sys.exit(-1)
    else:
        print('Подключение прошло успешно')

    app = QtWidgets.QApplication(sys.argv)

    MainForm = QtWidgets.QMainWindow()
    Ui_Main = Ui_MainForm()
    Ui_Main.setupUi(MainForm)

    Ui_Main.tableView_Activ.setModel(get_Activ())
    Ui_Main.tableView_act_isp.setModel(get_act_isp())

    Ui_Main.delete_button.clicked.connect(Delete_Dialog)
    Ui_Main.edit_button.clicked.connect(Edit_Dialog)
    Ui_Main.add_button.clicked.connect(Add_Dialog)
    Ui_Main.filter_button.clicked.connect(set_filter)
    Ui_Main.reset_button.clicked.connect(reset_filter)
    Ui_Main.exit_button.clicked.connect(MainForm.close)
    Ui_Main.analysis_button.clicked.connect(Analysis_Dialog)
    Ui_Main.save_button.clicked.connect(create_filter_doc)

    # выравнивание по ширине текста в столбце
    Ui_Main.tableView_Activ.resizeColumnsToContents()
    Ui_Main.tableView_act_isp.resizeColumnsToContents()

    #сортировка по столбцам
    Ui_Main.tableView_Activ.setSortingEnabled(True)
    Ui_Main.tableView_act_isp.setSortingEnabled(True)

    MainForm.show()
    sys.exit(app.exec())


